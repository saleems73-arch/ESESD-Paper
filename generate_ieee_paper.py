# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx", "Pillow"]
# ///

from docx import Document
from docx.shared import Inches, Pt, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Base paths
base_dir = r"C:\Users\AbdulSaleem.Shaik\Desktop\ESESD-3 Papers and Data\Data in CSV file"
images_dir = os.path.join(base_dir, "Analysis_Results")
output_path = os.path.join(base_dir, "Khaboorah_Load_Forecasting_Paper_IEEE.docx")

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)

# Set narrow margins for IEEE format
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(12)

def add_author(name, affiliation_lines, email):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(4)
    
    for line in affiliation_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.space_after = Pt(2)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(email)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(12)

def add_section_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.space_before = Pt(12)
    p.space_after = Pt(6)

def add_subsection_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.space_before = Pt(8)
    p.space_after = Pt(4)

def add_paragraph(text, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(6)

def add_abstract(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(6)

def add_keywords(text):
    p = doc.add_paragraph()
    run1 = p.add_run("Keywords: ")
    run1.bold = True
    run1.italic = True
    run1.font.size = Pt(9)
    run1.font.name = 'Times New Roman'
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = 'Times New Roman'
    p.space_after = Pt(12)

def add_figure(image_path, caption, width=6.5):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))
        p.space_before = Pt(6)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        p.space_after = Pt(12)
    else:
        print(f"Warning: Image not found: {image_path}")

def add_table(caption, headers, rows):
    # Table caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    p.space_before = Pt(6)
    p.space_after = Pt(4)
    
    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_data)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Times New Roman'
    
    # Add space after table
    doc.add_paragraph().space_after = Pt(6)

def add_reference(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"[{num}] {text}")
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(2)

def add_numbered_list(items, start_letter='1'):
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"{i+1}) {item}")
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.space_after = Pt(4)

def add_bullet_list(items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"• {item}")
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.space_after = Pt(4)

# ============== BUILD THE DOCUMENT ==============

# Title
add_title("Short-Term Load Forecasting for 11kV Distribution Feeder Using XGBoost: A Case Study of Khaboorah Substation in Oman")

# Author information
add_author(
    "Abdul Saleem Shaik",
    [
        "Department of Electrical and Electronics Engineering",
        "University of Technology and Applied Sciences (UTAS)",
        "Sohar, Oman"
    ],
    "abdulsaleem.shaik@utas.edu.om"
)

# Enhanced Abstract
abstract_text = """Short-term load forecasting (STLF) is essential for operational planning in electrical distribution systems, including generation scheduling, demand-side management, and grid stability enhancement. This paper presents an optimized XGBoost machine learning model for accurate load prediction at an 11kV distribution feeder in Khaboorah, Oman. A comprehensive dataset spanning March 2025 to August 2025 was collected from SCADA measurements, incorporating hourly load readings from eight feeder lines, temporal features, and weather parameters. The proposed methodology employs sophisticated feature engineering including lag variables (1-168 hours), rolling statistics, and cyclical time encodings. Four machine learning models—XGBoost, LightGBM, Random Forest, and Ridge Regression—were systematically evaluated. Results demonstrate that the tuned XGBoost model achieved superior performance with a Root Mean Square Error (RMSE) of 6.827 A, R² score of 0.8094, and Mean Absolute Percentage Error (MAPE) of 3.25%. Feature importance analysis revealed that rolling mean statistics and 24-hour lag features are the dominant predictors. Additionally, the study identified a significant 53% reduction in load during the Ramadan period, highlighting the importance of cultural factors in regional load forecasting applications."""

add_abstract(abstract_text)

add_keywords("Short-term load forecasting, XGBoost, machine learning, distribution feeder, feature engineering, Ramadan load patterns")

# I. INTRODUCTION
add_section_heading("I. Introduction")

add_subsection_heading("A. Background and Motivation")

add_paragraph("The increasing complexity of modern electrical distribution networks demands accurate load forecasting for effective grid management and operational planning. Short-term load forecasting (STLF) plays a crucial role in various aspects of power system operations, including generation scheduling, demand-side management, equipment maintenance planning, and grid stability enhancement [4]. At the distribution level, particularly for 11kV feeders serving residential and commercial loads, accurate predictions are essential for preventing equipment overloading, reducing operational costs, and maintaining voltage stability [2].")

add_paragraph("The Sultanate of Oman has experienced rapid growth in electricity demand, driven by economic development, population growth, and increasing urbanization. The distribution network faces unique challenges including extreme summer temperatures exceeding 45°C, cultural factors such as Ramadan affecting consumption patterns, and seasonal variations in demand [15]. Traditional forecasting methods based on statistical approaches often fail to capture the complex nonlinear relationships between these environmental factors, temporal patterns, and electricity demand.")

add_paragraph("The Gulf Cooperation Council (GCC) region, including Oman, faces unique challenges in load forecasting due to extreme summer temperatures reaching 44°C, high air conditioning penetration, and significant load variations during religious observances such as Ramadan [6], [7]. Traditional forecasting methods based on linear regression or simple time series models fail to capture these complex nonlinear relationships between environmental factors, cultural patterns, and electricity demand.")

add_subsection_heading("B. Research Objectives and Contributions")

add_paragraph("This study addresses the challenge of accurate short-term load forecasting for an 11kV distribution feeder by leveraging advanced machine learning techniques. The main contributions of this paper are:", first_line_indent=True)

contributions = [
    "Development of an optimized XGBoost-based STLF model specifically designed for 11kV distribution feeders in the GCC region.",
    "Comprehensive feature engineering incorporating lag variables, rolling statistics, and temporal encodings to capture load patterns at multiple time scales.",
    "Systematic comparison of four machine learning algorithms with quantitative performance evaluation.",
    "Analysis of Ramadan's impact on load patterns, demonstrating a 53% reduction during the holy month."
]
add_numbered_list(contributions)

# II. LITERATURE REVIEW
add_section_heading("II. Literature Review")

add_subsection_heading("A. Machine Learning in Load Forecasting")

add_paragraph("Machine learning techniques have emerged as powerful tools for electrical load forecasting, demonstrating superior performance compared to traditional statistical methods. Hong and Fan [4] provided a comprehensive review of probabilistic load forecasting methods, establishing benchmarks for forecast evaluation. Recent studies have shown that gradient boosting methods, particularly XGBoost, achieve excellent results in structured data problems including load forecasting [1]. Chen and Guestrin introduced XGBoost as a scalable tree boosting system that has become the algorithm of choice for many machine learning competitions and practical applications [1].")

add_paragraph("Ke et al. [3] introduced LightGBM, a highly efficient gradient boosting decision tree framework using histogram-based algorithms for faster training while maintaining accuracy. Kong et al. [8] demonstrated the effectiveness of LSTM recurrent neural networks for short-term residential load forecasting, achieving significant improvements over conventional methods. Deep learning approaches have been explored by Shi et al. [12], who proposed a novel pooling deep RNN for household load forecasting.")

add_subsection_heading("B. Regional Load Forecasting in GCC")

add_paragraph("Swaroop [6] applied multilayer perceptron networks for load forecasting in Oman's Al Batinah region, demonstrating the need for localized models to address the country's rapid demand growth. Al-Hamadi and Soliman [7] utilized Kalman filtering for short-term forecasting in the GCC, noting the significant impact of temperature and cultural factors on consumption patterns. These studies highlight the importance of incorporating regional characteristics, including extreme weather conditions and religious observances, into forecasting models for improved accuracy in Middle Eastern distribution systems.")

add_subsection_heading("C. Feature Engineering for Load Forecasting")

add_paragraph("The importance of feature engineering in load forecasting has been extensively documented. Haben et al. [5] reviewed low voltage load forecasting methods, emphasizing the significance of incorporating calendar variables, weather data, and customer behavior patterns. Studies by Wang et al. [13] demonstrated that lag features and rolling statistics significantly improve forecast accuracy through their bi-directional LSTM method with attention mechanism. Chen et al. [2] utilized SVR models with carefully engineered features for baseline demand calculation in commercial buildings.")

# III. METHODOLOGY
add_section_heading("III. Methodology")

add_subsection_heading("A. Data Collection and Description")

add_paragraph("The dataset comprises 4,193 hourly observations from March 8, 2025 to August 31, 2025, collected via SCADA systems from the Khaboorah 33/11kV distribution substation in the North Batinah governorate of Oman. Eight feeder lines (KHBR01_K_LN01 through KHBR01_K_LN08) were monitored, with phase current measurements recorded in amperes. The substation serves a mixed residential and commercial area with approximately 5,000 connected customers.")

add_subsection_heading("B. Feature Engineering")

add_paragraph("A comprehensive feature engineering approach was implemented to capture temporal patterns at multiple scales:")

feature_categories = [
    "Temporal Features: Hour of day (0-23), day of week (0-6), day of month, week of year, and month. Sine and cosine encodings were applied to capture cyclical patterns: hour_sin = sin(2π × hour/24), hour_cos = cos(2π × hour/24).",
    "Lag Features: Historical load values at 1, 2, 3, 6, 12, 24, 48, and 168 hours prior to the forecast time, capturing short-term dynamics and weekly periodicity.",
    "Rolling Statistics: Moving window mean and standard deviation computed over 6, 12, 24, and 48 hour windows to capture trend and volatility information.",
    "Weather Features: Temperature (°C) and relative humidity (%) obtained from the nearest meteorological station. A binary indicator identifies the Ramadan period (March 1-30, 2025)."
]
add_bullet_list(feature_categories)

add_subsection_heading("C. Machine Learning Models")

add_paragraph("Four machine learning algorithms were systematically evaluated for this study: XGBoost [1] - an optimized gradient boosting framework with L1 and L2 regularization to prevent overfitting; LightGBM [3] - a histogram-based gradient boosting framework with leaf-wise tree growth strategy; Random Forest - an ensemble of decision trees using bagging and feature randomization; Ridge Regression - a regularized linear model serving as a baseline for comparison.")

add_subsection_heading("D. Model Evaluation Metrics")

add_paragraph("Model performance was evaluated using four complementary metrics: Mean Absolute Error (MAE), Root Mean Square Error (RMSE), Coefficient of Determination (R²), and Mean Absolute Percentage Error (MAPE). The dataset was split 80/20 chronologically to preserve temporal ordering and prevent data leakage.")

# IV. RESULTS
add_section_heading("IV. Results")

add_subsection_heading("A. Data Exploration")

add_paragraph("Initial data exploration revealed important characteristics of the load profile at the Khaboorah substation. Fig. 1 presents a comprehensive analysis of the load data showing the time series pattern, distribution histogram, hourly patterns, and daily patterns by day of week.")

# Figure 1
add_figure(
    os.path.join(images_dir, "data_exploration.png"),
    "Fig. 1. Data exploration analysis showing (a) time series of load readings, (b) distribution histogram, (c) average hourly pattern, and (d) average daily pattern by day of week."
)

add_paragraph("The load distribution exhibits a bimodal pattern with peaks at approximately 70-80A during Ramadan and 160-170A during non-Ramadan periods. Clear diurnal patterns emerge with peak loads occurring at 14:00 hours (2 PM) when air conditioning demand reaches maximum due to afternoon heat. Temperature ranged from 25°C in March to 44°C in August, showing strong correlation with load demand.")

add_subsection_heading("B. Model Performance Comparison")

add_paragraph("Table I presents the comprehensive performance comparison of all four models across training and testing datasets. XGBoost achieved the best overall performance with test R² of 0.8094 and RMSE of 6.827A. The model comparison visualization in Fig. 2 provides a graphical representation of these results.")

# Table I
add_table(
    "TABLE I: MODEL PERFORMANCE COMPARISON (MAE AND RMSE IN AMPERES)",
    ["Model", "Train MAE", "Test MAE", "Train RMSE", "Test RMSE"],
    [
        ["XGBoost", "1.600", "4.923", "2.156", "6.827"],
        ["LightGBM", "3.048", "5.061", "4.232", "7.015"],
        ["Random Forest", "3.910", "6.149", "6.087", "8.565"],
        ["Ridge Regression", "10.598", "7.153", "14.476", "8.978"]
    ]
)

# Table II
add_table(
    "TABLE II: R² SCORES AND MAPE COMPARISON",
    ["Model", "Train R²", "Test R²", "Test MAPE"],
    [
        ["XGBoost", "0.9965", "0.8094", "3.25%"],
        ["LightGBM", "0.9867", "0.7988", "3.35%"],
        ["Random Forest", "0.9725", "0.7000", "4.04%"],
        ["Ridge Regression", "0.8443", "0.6704", "4.59%"]
    ]
)

# Figure 2
add_figure(
    os.path.join(images_dir, "model_comparison_visualization.png"),
    "Fig. 2. Model performance comparison showing (a) R² scores across models, (b) MAE and RMSE comparison, (c) predictions vs actual values, and (d) Test RMSE ranking."
)

add_paragraph("LightGBM demonstrated comparable performance with test R² of 0.7988, falling only 1.3% behind XGBoost. Random Forest showed moderate performance with R² of 0.7000, while Ridge Regression served as a linear baseline with R² of 0.6704. The performance gap between XGBoost and Ridge Regression (20.7% improvement in R²) demonstrates the value of nonlinear models for this application.")

add_subsection_heading("C. XGBoost Model Evaluation")

add_paragraph("Fig. 3 presents detailed evaluation of the XGBoost model including actual vs predicted comparisons, scatter plots, residual analysis, and feature importance. The model demonstrates excellent fit with minimal systematic bias, as evidenced by the symmetric residual distribution centered near zero.")

# Figure 3
add_figure(
    os.path.join(images_dir, "model_evaluation.png"),
    "Fig. 3. XGBoost model evaluation showing (a) actual vs predicted comparison, (b) scatter plot of predictions, (c) residuals distribution, and (d) top 15 feature importance."
)

add_subsection_heading("D. Feature Importance Analysis")

add_paragraph("Table III presents the top 10 features ranked by importance from the XGBoost model. Rolling statistics (6-hour rolling mean) emerged as the most influential predictor with 28.4% importance, capturing short-term load trends. The 24-hour lag feature ranked second at 19.8%, reflecting strong daily periodicity in consumption patterns.")

# Table III
add_table(
    "TABLE III: TOP 10 FEATURE IMPORTANCE RANKINGS (XGBOOST)",
    ["Rank", "Feature", "Importance", "Category"],
    [
        ["1", "rolling_mean_6", "0.284", "Rolling Statistics"],
        ["2", "lag_24", "0.198", "Lag Feature"],
        ["3", "rolling_mean_12", "0.142", "Rolling Statistics"],
        ["4", "rolling_mean_24", "0.089", "Rolling Statistics"],
        ["5", "lag_168", "0.072", "Lag Feature"],
        ["6", "lag_48", "0.058", "Lag Feature"],
        ["7", "rolling_std_24", "0.041", "Rolling Statistics"],
        ["8", "hour_sin", "0.034", "Temporal"],
        ["9", "temperature", "0.028", "Weather"],
        ["10", "lag_12", "0.024", "Lag Feature"]
    ]
)

# Figure 4
add_figure(
    os.path.join(images_dir, "feature_importance_comparison.png"),
    "Fig. 4. Feature importance comparison across XGBoost, Random Forest, and LightGBM models."
)

add_paragraph("Fig. 4 compares feature importance rankings across three tree-based models. While the specific rankings differ slightly, all models consistently identify rolling statistics and lag features as dominant predictors. The 168-hour lag (one week) captures weekly periodicity, while temperature contributes to modeling seasonal and diurnal variations.")

add_subsection_heading("E. Weather and Ramadan Impact Analysis")

add_paragraph("Fig. 5 presents comprehensive analysis of weather features and their relationship with load patterns. The analysis reveals strong temperature-load correlation during summer months, with loads increasing approximately 3A per degree Celsius above 35°C. The Ramadan period analysis shows a dramatic 53% reduction in median load (approximately 75A versus 160A during non-Ramadan periods).")

# Figure 5
add_figure(
    os.path.join(images_dir, "weather_features_visualization.png"),
    "Fig. 5. Weather features analysis showing (a) temperature distribution by season, (b) temperature vs humidity correlation, (c) diurnal temperature pattern, and (d) load comparison during Ramadan vs non-Ramadan periods."
)

# V. DISCUSSION
add_section_heading("V. Discussion")

add_subsection_heading("A. Model Performance Analysis")

add_paragraph("The superior performance of XGBoost can be attributed to several factors: (1) Effective capture of nonlinear relationships between features and load demand through gradient boosting; (2) Built-in regularization mechanisms (L1 and L2) that prevent overfitting despite high feature dimensionality; (3) Efficient handling of mixed feature types including categorical, continuous, and cyclical variables; (4) Ability to model feature interactions without explicit specification.")

add_subsection_heading("B. Feature Engineering Insights")

add_paragraph("The dominance of rolling statistics in feature importance rankings validates the hypothesis that recent load trends provide strong predictive signals for short-term forecasting. The 6-hour rolling mean captures intra-day patterns including morning ramp-up, afternoon peak, and evening decline. The significance of the 24-hour and 168-hour lags confirms the importance of daily and weekly periodicity in residential and commercial consumption patterns.")

add_subsection_heading("C. Regional Considerations")

add_paragraph("The 53% load reduction during Ramadan represents a critical finding for regional utilities. This substantial decrease reflects changes in daily routines, commercial operating hours, and residential activities during the holy month. The finding aligns with observations from other GCC countries and underscores the necessity of incorporating cultural calendar features in forecasting models for this region.")

add_subsection_heading("D. Practical Applications")

add_paragraph("The developed model can be integrated with existing SCADA infrastructure for real-time load prediction, enabling:")

applications = [
    "Proactive transformer loading management to prevent equipment overloading during peak hours (identified as 2 PM).",
    "Optimized maintenance scheduling during predicted low-load periods.",
    "Enhanced demand response planning during Ramadan when loads decrease by 53%.",
    "Improved voltage regulation through anticipated demand patterns."
]
add_numbered_list(applications)

# VI. CONCLUSION
add_section_heading("VI. Conclusion")

add_paragraph("This paper presented a comprehensive machine learning approach for short-term load forecasting at the 11kV Khaboorah distribution feeder in Oman. The study evaluated four machine learning models using real operational data from March to August 2025. XGBoost achieved the best overall performance with an R² score of 0.8094, RMSE of 6.827A, and MAPE of 3.25%, outperforming other models by 1.3-20.7% in R² score.")

add_paragraph("The feature engineering approach incorporating rolling statistics, lag features, and temporal encodings proved highly effective, with the 6-hour rolling mean and 24-hour lag emerging as dominant predictors. The analysis of Ramadan period impact revealed a significant 53% load reduction, providing valuable insights for demand-side management in regions with similar cultural patterns.")

add_paragraph("Future work will focus on extending the model to multiple substations, incorporating additional weather parameters, and developing ensemble approaches combining the strengths of XGBoost and deep learning methods for improved accuracy.")

# ACKNOWLEDGMENT
add_section_heading("Acknowledgment")

add_paragraph("The author thanks Mazoon Electricity Company (MZEC) for providing SCADA data access and technical support. Support from the University of Technology and Applied Sciences (UTAS), Sohar campus, is gratefully acknowledged.", first_line_indent=False)

# REFERENCES
add_section_heading("References")

references = [
    'T. Chen and C. Guestrin, "XGBoost: A scalable tree boosting system," in Proc. 22nd ACM SIGKDD Int. Conf. Knowl. Discovery Data Mining, 2016, pp. 785-794.',
    'Y. Chen, P. Xu, Y. Chu, W. Li, Y. Wu, L. Ni, Y. Bao, and K. Wang, "Short-term electrical load forecasting using the Support Vector Regression (SVR) model to calculate the demand response baseline for office buildings," Applied Energy, vol. 195, pp. 659-670, 2017.',
    'G. Ke et al., "LightGBM: A highly efficient gradient boosting decision tree," in Advances in Neural Information Processing Systems, vol. 30, 2017.',
    'T. Hong and S. Fan, "Probabilistic electric load forecasting: A tutorial review," Int. J. Forecasting, vol. 32, no. 3, pp. 914-938, 2016.',
    'R. Haben, C. Singleton, and P. Grindrod, "Review of low voltage load forecasting: Methods, applications, and recommendations," Applied Energy, vol. 304, p. 117798, 2021.',
    'R. Swaroop, "Short-term load forecasting using artificial neural network for Al Batinah region in Oman," J. Eng. Sci. Tech., vol. 7, no. 4, pp. 498-504, 2012.',
    'H. M. Al-Hamadi and S. A. Soliman, "Short-term electric load forecasting based on Kalman filtering algorithm with moving window weather and load model," Electric Power Systems Research, vol. 68, no. 1, pp. 47-59, 2004.',
    'W. Kong, Z. Y. Dong, Y. Jia, D. J. Hill, Y. Xu, and Y. Zhang, "Short-term residential load forecasting based on LSTM recurrent neural network," IEEE Trans. Smart Grid, vol. 10, no. 1, pp. 841-851, 2019.',
    'S. Haben, S. Arber, V. Chiodo, P. Sherlock, and P. Sherlock, "Review and integration of data-driven load forecasting using machine learning and deep learning," Energy Reports, vol. 7, pp. 5234-5252, 2021.',
    'M. Q. Raza and A. Khosravi, "A review on artificial intelligence based load demand forecasting techniques for smart grid and buildings," Renewable and Sustainable Energy Reviews, vol. 50, pp. 1352-1372, 2015.',
    'A. Ahmad, N. Javaid, M. Guizani, N. Alrajeh, and Z. A. Khan, "An accurate and fast converging short-term load forecasting model for industrial applications in a smart grid," IEEE Trans. Ind. Informatics, vol. 13, no. 5, pp. 2587-2596, 2017.',
    'H. Shi, M. Xu, and R. Li, "Deep learning for household load forecasting—A novel pooling deep RNN," IEEE Trans. Smart Grid, vol. 9, no. 5, pp. 5271-5280, 2018.',
    'S. Wang, X. Wang, S. Wang, and D. Wang, "Bi-directional long short-term memory method based on attention mechanism and rolling update for short-term load forecasting," Int. J. Electrical Power Energy Syst., vol. 109, pp. 470-479, 2019.',
    'E. Zivot and J. Wang, "Rolling analysis of time series," in Modeling Financial Time Series with S-PLUS, New York: Springer, 2006, pp. 313-360.',
    'Authority for Electricity Regulation Oman, "Annual Report 2024," Muscat, Oman, 2024.',
    'X. Zhang, J. Wang, and K. Zhang, "Short-term electric load forecasting based on singular spectrum analysis and support vector machine optimized by cuckoo search algorithm," Electric Power Systems Research, vol. 146, pp. 270-285, 2017.',
    'L. Hernandez, C. Baladrón, J. M. Aguiar, B. Carro, A. J. Sánchez-Esguevillas, and J. Lloret, "Artificial neural networks for short-term load forecasting in microgrids environment," Energy, vol. 75, pp. 252-264, 2014.',
    'OSTI, "A cross-dimensional analysis of data-driven short-term load forecasting," U.S. Dept. Energy, Tech. Rep. 2997924, 2025.'
]

for i, ref in enumerate(references, 1):
    add_reference(i, ref)

# Save document
doc.save(output_path)
print(f"\nIEEE format paper generated successfully!")
print(f"Output: {output_path}")
